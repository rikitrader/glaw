#!/usr/bin/env python3
"""GLAW md -> Word .docx (legal formatting). Renders a [BRACKET] markdown master into a real Microsoft Word
document in the us-corporate-legal-instrument house style: US Letter, margins 1in (1.14in bottom), Times New Roman
10pt justified body (x1.15 leading), centered ALL-CAPS bold title/ARTICLE headings, bold run-in Section/(a) leads,
no first-line indent, bordered tables, small italic disclaimer notes. Brackets preserved (blank fillable).

Usage: md_to_docx.py <in.md> <out.docx>   (or)   md_to_docx.py --dir <srcdir> <outdir>
"""
import sys, os, re, glob
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def setfont(run, size=10, bold=False, italic=False, color=(0,0,0)):
    run.font.name = "Times New Roman"; run.font.size = Pt(size)
    run.bold = bold; run.italic = italic; run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr(); rf = rpr.get_or_add_rFonts()
    for a in ("w:ascii","w:hAnsi","w:cs"): rf.set(qn(a), "Times New Roman")

def add_inline(p, text, base_bold=False, size=10):
    # split on **bold**
    for i, seg in enumerate(re.split(r"(\*\*[^*]+\*\*)", text)):
        if not seg: continue
        if seg.startswith("**") and seg.endswith("**"):
            setfont(p.add_run(seg[2:-2]), size=size, bold=True)
        else:
            setfont(p.add_run(seg), size=size, bold=base_bold)

def cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr(); b = OxmlElement('w:tcBorders')
    for edge in ("top","left","bottom","right"):
        e = OxmlElement(f'w:{edge}'); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'4'); e.set(qn('w:color'),'000000'); b.append(e)
    tcPr.append(b)

def convert(src, out):
    md = open(src).read().split("\n")
    doc = Document()
    s = doc.sections[0]
    s.page_width, s.page_height = Inches(8.5), Inches(11)
    s.top_margin = Inches(1); s.bottom_margin = Inches(1.14); s.left_margin = Inches(1); s.right_margin = Inches(0.98)
    st = doc.styles["Normal"]; st.font.name = "Times New Roman"; st.font.size = Pt(10)
    st.paragraph_format.line_spacing = 1.15; st.paragraph_format.space_after = Pt(6)
    i = 0
    while i < len(md):
        ln = md[i]
        if not ln.strip(): i += 1; continue
        if re.match(r"^---+$", ln):
            p = doc.add_paragraph(); pb = OxmlElement('w:pBdr'); bt = OxmlElement('w:bottom')
            bt.set(qn('w:val'),'single'); bt.set(qn('w:sz'),'6'); bt.set(qn('w:color'),'000000'); pb.append(bt)
            p._p.get_or_add_pPr().append(pb); i += 1; continue
        m = re.match(r"^(#{1,4})\s+(.*)", ln)
        if m:
            lvl, txt = len(m.group(1)), m.group(2)
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12 if lvl<=2 else 8)
            if lvl == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER; setfont(p.add_run(txt), 10, bold=True)
            elif lvl == 2:                       # ARTICLE: centered, ALL CAPS, bold, no underline
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER; setfont(p.add_run(txt.upper()), 10, bold=True)
            else:                                # subsection heading line: bold, flush-left, no underline
                setfont(p.add_run(txt), 10, bold=True)
            i += 1; continue
        if ln.startswith(">"):
            buf = []
            while i < len(md) and md[i].startswith(">"): buf.append(re.sub(r"^>\s?","",md[i])); i += 1
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8)
            setfont(p.add_run(" ".join(buf)), 9, italic=True, color=(80,80,80)); continue
        if "|" in ln and i+1 < len(md) and re.match(r"^\s*\|?[\s:|-]+\|?\s*$", md[i+1]) and "-" in md[i+1]:
            cells = lambda r: [c.strip() for c in r.strip().strip("|").split("|")]
            hdr = cells(ln); i += 2; rows = []
            while i < len(md) and "|" in md[i] and md[i].strip(): rows.append(cells(md[i])); i += 1
            tbl = doc.add_table(rows=1, cols=len(hdr))
            for j, h in enumerate(hdr):
                c = tbl.rows[0].cells[j]; c.paragraphs[0].text = ""; add_inline(c.paragraphs[0], h, base_bold=True, size=10); cell_borders(c)
            for r in rows:
                rc = tbl.add_row().cells
                for j in range(len(hdr)):
                    c = rc[j]; c.paragraphs[0].text = ""; add_inline(c.paragraphs[0], r[j] if j < len(r) else "", size=10); cell_borders(c)
            doc.add_paragraph(); continue
        if re.match(r"^\s*([-*]|\d+\.)\s+", ln):
            while i < len(md) and re.match(r"^\s*([-*]|\d+\.)\s+", md[i]):
                txt = re.sub(r"^\s*([-*]|\d+\.)\s+","",md[i])
                p = doc.add_paragraph(style="List Bullet" if md[i].lstrip()[0] in "-*" else "List Number")
                add_inline(p, txt); i += 1
            continue
        # paragraph (justified, first-line indent for body prose)
        buf = []
        while i < len(md) and md[i].strip() and not re.match(r"^(#|>|\||---|\s*[-*]|\s*\d+\.)", md[i]): buf.append(md[i]); i += 1
        if not buf:                      # no-progress guard: a stray "|" line etc. -> emit it + advance
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; add_inline(p, md[i]); i += 1; continue
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_inline(p, " ".join(buf))
    doc.save(out)
    return out

def main():
    if "--dir" in sys.argv:
        idx = sys.argv.index("--dir"); src = sys.argv[idx+1]; outdir = sys.argv[idx+2]
        os.makedirs(outdir, exist_ok=True); n = 0
        for f in sorted(glob.glob(os.path.join(src, "*.md"))):
            base = os.path.splitext(os.path.basename(f))[0]
            convert(f, os.path.join(outdir, base + ".docx")); n += 1
        print(f"wrote {n} .docx to {outdir}")
    else:
        convert(sys.argv[1], sys.argv[2]); print("wrote", sys.argv[2])

if __name__ == "__main__":
    main()
