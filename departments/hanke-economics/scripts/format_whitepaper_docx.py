#!/usr/bin/env python3
"""Apply formal white-paper presentation to an already converted DOCX."""
from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_run_font(run, size=10.5, bold=None, italic=None, color=None):
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), "Aptos")


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def main(src, dst):
    doc = Document(src)
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in (("Heading 1", 15, (31, 78, 121)), ("Heading 2", 12.5, (31, 78, 121)), ("Heading 3", 11, (55, 55, 55))):
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(*color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(5)

    # The converter emits direct-formatted headings; normalize them by position.
    heading_count = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if text == "[[PAGEBREAK]]":
            paragraph.clear()
            paragraph.add_run().add_break(WD_BREAK.PAGE)
            continue
        for run in paragraph.runs:
            set_run_font(run, 10.5)
        if text.startswith("Venezuela Dollarization White Paper"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(10)
            paragraph.paragraph_format.space_after = Pt(4)
            for run in paragraph.runs:
                set_run_font(run, 22, bold=True, color=(31, 78, 121))
        elif text.startswith("Monetary, Banking,"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, 12, italic=True, color=(80, 80, 80))
        elif text.startswith("#"):
            # Defensive handling if an unconverted markdown heading remains.
            heading_count += 1
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_run_font(run, 14 if heading_count == 1 else 11.5, bold=True, color=(31, 78, 121))
        elif text.startswith("PART "):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(16)
            paragraph.paragraph_format.space_after = Pt(7)
            for run in paragraph.runs:
                set_run_font(run, 15, bold=True, color=(31, 78, 121))
        elif text[:2].isdigit() or text.startswith("Evidence labels") or text.startswith("Source-of-truth"):
            paragraph.paragraph_format.space_before = Pt(9)
            for run in paragraph.runs:
                set_run_font(run, 11.5, bold=True, color=(55, 55, 55))

    # Tables: readable professional style and repeated header rows.
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for idx, row in enumerate(table.rows):
            if idx == 0:
                tr_pr = row._tr.get_or_add_trPr()
                tbl_header = OxmlElement("w:tblHeader")
                tbl_header.set(qn("w:val"), "true")
                tr_pr.append(tbl_header)
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(2)
                    for run in p.runs:
                        set_run_font(run, 9.2, bold=(idx == 0))
                if idx == 0:
                    shade_cell(cell, "D9EAF7")

    # Header/footer identify the edition and add a page field.
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.text = "HAEIS  |  Venezuela Dollarization White Paper  |  Preliminary Review Edition"
    for run in header.runs:
        set_run_font(run, 8, color=(100, 100, 100))
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "UNVERIFIED / NON-DECISION-GRADE  •  Page "
    for run in footer.runs:
        set_run_font(run, 8, color=(100, 100, 100))
    add_field(footer, "PAGE")
    for run in footer.runs:
        set_run_font(run, 8, color=(100, 100, 100))

    core = doc.core_properties
    core.title = "Venezuela Dollarization White Paper — Preliminary Unverified Review"
    core.subject = "HAEIS sovereign-macro analysis of Venezuela monetary-regime options"
    core.author = "Hanke Applied Economics Intelligence System (HAEIS)"
    core.comments = "Preliminary reading draft. Not decision-grade. Missing and disputed data are explicitly marked."
    doc.save(dst)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("usage: format_whitepaper_docx.py INPUT.docx OUTPUT.docx")
    main(Path(sys.argv[1]), Path(sys.argv[2]))
