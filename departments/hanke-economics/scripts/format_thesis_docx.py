#!/usr/bin/env python3
"""Render the HAEIS thesis markdown as a formal, double-spaced thesis DOCX."""
from pathlib import Path
import sys

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from format_whitepaper_docx import main as base_format


def add_page_field(paragraph):
    paragraph.text = "HAEIS THESIS REVIEW EDITION  •  Page "
    run = paragraph.add_run()
    for kind, value in (("begin", None), (None, "PAGE"), ("end", None)):
        if kind:
            node = OxmlElement("w:fldChar")
            node.set(qn("w:fldCharType"), kind)
        else:
            node = OxmlElement("w:instrText")
            node.set(qn("xml:space"), "preserve")
            node.text = value
        run._r.append(node)


def main(src, dst):
    base = Path(dst).with_suffix(".base.docx")
    base_format(Path(src), base)
    doc = Document(base)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_after = Pt(0)
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "[[PAGEBREAK]]":
            paragraph.clear()
            paragraph.add_run().add_break(WD_BREAK.PAGE)
            continue
        if paragraph.text.strip():
            paragraph.paragraph_format.line_spacing = 2.0
            if not paragraph.text.startswith("PART ") and not paragraph.text.startswith("Venezuela Dollarization"):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(10)
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(footer)
    doc.core_properties.title = "Venezuela Dollarization — Thesis Research Report"
    doc.core_properties.subject = "HAEIS evidence-led sovereign-macro thesis review edition"
    doc.core_properties.author = "Hanke Applied Economics Intelligence System (HAEIS)"
    doc.save(dst)
    base.unlink(missing_ok=True)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("usage: format_thesis_docx.py INPUT.docx OUTPUT.docx")
    main(Path(sys.argv[1]), Path(sys.argv[2]))
