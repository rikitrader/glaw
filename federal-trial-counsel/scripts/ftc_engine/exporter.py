"""
Document Exporter — Converts FTC templates and drafts to court-formatted .docx files.

Produces Word documents with proper federal court formatting:
- Times New Roman 12pt
- Double-spaced body text
- 1-inch margins
- Centered headings (ALL CAPS, bold)
- Right-aligned signature blocks
- Numbered paragraphs with proper indentation

Output .docx files open directly in Microsoft Word, Google Docs, or LibreOffice.
Export to PDF from any of those applications.

Usage:
  ftc export --draft --case case.json --out complaint.docx
  ftc export --template motions/motion_to_dismiss --case case.json --out motion.docx
  ftc export --text input.md --out output.docx
"""
from __future__ import annotations

import re
from pathlib import Path
from dataclasses import dataclass, field

from docx import Document
from docx.shared import Inches, Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT


# ── Court formatting constants ───────────────────────────────────────────────

FONT_NAME = "Times New Roman"
FONT_SIZE_BODY = Pt(12)
FONT_SIZE_HEADING = Pt(12)
FONT_SIZE_CAPTION = Pt(12)
LINE_SPACING = 2.0  # Double-spaced
MARGIN = Inches(1)
FIRST_LINE_INDENT = Inches(0.5)

# Patterns for parsing legal document structure
RE_HEADING = re.compile(r"^\s{5,}([A-Z][A-Z .,'&§()/:;0-9—–-]{3,})\s*$")
RE_SUBHEADING = re.compile(r"^\s{5,}([A-Z]\.?\s+.+)$")
RE_NUMBERED = re.compile(r"^\s{3,}(\d+)\.\s+(.+)")
RE_LETTERED = re.compile(r"^\s{6,}(\[?\s*[a-z]\s*\]?)\s*[.)]\s*(.+)")
RE_CHECKBOX = re.compile(r"^\s{3,}\[\s*\]\s+(\d+)\.\s+(.+)")
RE_SIGNATURE = re.compile(r"^\s{20,}(Respectfully submitted|/s/|Copies to:|DONE AND ORDERED|I HEREBY CERTIFY)")
RE_DASHES = re.compile(r"^_{5,}|^-{5,}")
RE_CAPTION_PARTY = re.compile(r"^\s*(.*?),\s*$")


@dataclass
class ExportResult:
    output_path: str
    format: str  # "docx"
    pages_estimate: int
    sections: int


def _setup_document(district_code: str | None = None) -> Document:
    """Create a new document with court-standard formatting.

    Args:
        district_code: Optional district code for district-specific formatting.
                       Falls back to module-level constants if None.
    """
    doc = Document()

    # Get formatting config from district if available
    font_name = FONT_NAME
    font_size = FONT_SIZE_BODY
    line_spacing = LINE_SPACING
    margin = MARGIN

    if district_code:
        try:
            from .districts import get_formatting_config
            fmt = get_formatting_config(district_code)
            font_name = fmt.get("font_name", FONT_NAME)
            font_size = Pt(fmt.get("font_size_pt", 12))
            line_spacing = fmt.get("line_spacing", LINE_SPACING)
            margin = Inches(fmt.get("margin_inches", 1.0))
        except Exception:
            pass  # Fall back to defaults

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = font_name
    font.size = font_size

    pf = style.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    # Set margins
    for section in doc.sections:
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin

    return doc


def _add_centered(doc: Document, text: str, bold: bool = False, size=None, caps: bool = False):
    """Add centered paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text.upper() if caps else text)
    run.font.name = FONT_NAME
    run.font.size = size or FONT_SIZE_BODY
    run.bold = bold
    return p


def _add_right_aligned(doc: Document, text: str, bold: bool = False):
    """Add right-aligned paragraph (for signatures)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.line_spacing = 1.0  # Single-spaced signature blocks
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_BODY
    run.bold = bold
    return p


def _add_body(doc: Document, text: str, indent: bool = True, bold: bool = False):
    """Add body paragraph with optional first-line indent."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_BODY
    run.bold = bold
    return p


def _add_blank_line(doc: Document):
    """Add empty paragraph for spacing."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("")
    run.font.size = FONT_SIZE_BODY
    return p


def _is_caption_area(lines: list[str], i: int) -> bool:
    """Detect if we're in the case caption block."""
    # Look for "UNITED STATES DISTRICT COURT" nearby
    window = lines[max(0, i - 10):i + 5]
    return any("UNITED STATES DISTRICT COURT" in l for l in window)


def _parse_and_render(doc: Document, text: str):
    """Parse legal document text and render to Word document."""
    lines = text.split("\n")
    i = 0
    in_signature_block = False
    caption_done = False
    section_count = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines (we control spacing ourselves)
        if not stripped:
            i += 1
            continue

        # ── Caption block (UNITED STATES DISTRICT COURT) ──
        if "UNITED STATES DISTRICT COURT" in stripped and not caption_done:
            _add_centered(doc, "UNITED STATES DISTRICT COURT", bold=True)
            i += 1
            # Read next few lines as caption
            while i < len(lines):
                cl = lines[i].strip()
                if not cl:
                    i += 1
                    continue
                if cl.startswith("_") or cl.startswith("-"):
                    _add_centered(doc, "_" * 40)
                    caption_done = True
                    i += 1
                    break
                elif "v." in cl and "Case No." in cl:
                    # Split "v." and "Case No." onto the same line
                    _add_body(doc, cl, indent=False)
                    i += 1
                elif cl.startswith("v."):
                    # "v." line — check if case no is on same line
                    _add_body(doc, cl, indent=False)
                    i += 1
                elif "Case No." in cl:
                    _add_right_aligned(doc, cl)
                    i += 1
                elif cl.endswith(","):
                    # Party name line
                    _add_body(doc, cl, indent=False)
                    i += 1
                elif cl in ("Plaintiff,", "Defendant.", "Defendant,",
                            "Plaintiff-Appellant,", "Plaintiff-Appellee,",
                            "Defendant-Appellant,", "Defendant-Appellee,",
                            "Defendant(s).", "Third-Party Defendant.",
                            "Plaintiff / Counterclaim Defendant,",
                            "Defendant / Counterclaim Plaintiff.",
                            "Defendant / Third-Party Plaintiff,"):
                    _add_right_aligned(doc, cl)
                    i += 1
                else:
                    _add_centered(doc, cl, bold=("DISTRICT" in cl.upper() or "DIVISION" in cl.upper()))
                    i += 1
            _add_blank_line(doc)
            continue

        # ── Document title (after caption, centered, all caps) ──
        if caption_done and RE_HEADING.match(line) and not in_signature_block:
            heading_text = stripped
            # Check for multi-line headings
            while i + 1 < len(lines) and lines[i + 1].strip() and RE_HEADING.match(lines[i + 1]):
                i += 1
                heading_text += "\n" + lines[i].strip()
            _add_blank_line(doc)
            _add_centered(doc, heading_text, bold=True)
            _add_blank_line(doc)
            section_count += 1
            i += 1
            continue

        # ── Signature / right-aligned blocks ──
        if RE_SIGNATURE.match(line):
            in_signature_block = True

        if in_signature_block:
            text_content = stripped
            if text_content.startswith("/s/"):
                _add_right_aligned(doc, text_content)
            elif text_content.startswith("Respectfully") or text_content.startswith("DONE AND ORDERED"):
                _add_blank_line(doc)
                _add_body(doc, text_content, indent=True)
                if "DONE AND ORDERED" in text_content:
                    in_signature_block = False
            elif text_content.startswith("_"):
                _add_right_aligned(doc, text_content)
            elif text_content.startswith("I HEREBY CERTIFY"):
                _add_blank_line(doc)
                _add_body(doc, text_content, indent=True)
            elif text_content.startswith("Copies to:"):
                _add_blank_line(doc)
                _add_body(doc, text_content, indent=False)
            else:
                _add_right_aligned(doc, text_content)
            i += 1
            continue

        # ── Checkbox paragraphs [ ] ──
        m = RE_CHECKBOX.match(line)
        if m:
            num, content = m.group(1), m.group(2)
            # Collect continuation lines
            while i + 1 < len(lines) and lines[i + 1].strip() and not RE_CHECKBOX.match(lines[i + 1]) and not RE_NUMBERED.match(lines[i + 1]) and not RE_HEADING.match(lines[i + 1]):
                i += 1
                content += " " + lines[i].strip()
            _add_body(doc, f"[ ] {num}. {content}", indent=True)
            i += 1
            continue

        # ── Numbered paragraphs ──
        m = RE_NUMBERED.match(line)
        if m:
            num, content = m.group(1), m.group(2)
            # Collect continuation lines
            while i + 1 < len(lines) and lines[i + 1].strip() and not RE_NUMBERED.match(lines[i + 1]) and not RE_HEADING.match(lines[i + 1]) and not RE_CHECKBOX.match(lines[i + 1]) and not RE_LETTERED.match(lines[i + 1]) and not RE_SIGNATURE.match(lines[i + 1]):
                i += 1
                content += " " + lines[i].strip()
            _add_body(doc, f"{num}. {content}", indent=True)
            i += 1
            continue

        # ── Lettered sub-paragraphs ──
        m = RE_LETTERED.match(line)
        if m:
            letter, content = m.group(1), m.group(2)
            while i + 1 < len(lines) and lines[i + 1].strip() and not RE_LETTERED.match(lines[i + 1]) and not RE_NUMBERED.match(lines[i + 1]) and not RE_HEADING.match(lines[i + 1]):
                i += 1
                content += " " + lines[i].strip()
            p = _add_body(doc, f"{letter.strip()}. {content}", indent=False)
            p.paragraph_format.left_indent = Inches(1.0)
            i += 1
            continue

        # ── Divider lines ──
        if RE_DASHES.match(stripped):
            i += 1
            continue

        # ── Roman numeral subheadings (I. II. III. etc) ──
        if re.match(r"^[IVX]+\.\s+", stripped):
            _add_blank_line(doc)
            _add_centered(doc, stripped, bold=True)
            i += 1
            continue

        # ── Regular body text ──
        content = stripped
        # Collect continuation
        while i + 1 < len(lines) and lines[i + 1].strip() and not RE_NUMBERED.match(lines[i + 1]) and not RE_HEADING.match(lines[i + 1]) and not RE_CHECKBOX.match(lines[i + 1]) and not RE_SIGNATURE.match(lines[i + 1]):
            next_stripped = lines[i + 1].strip()
            if not next_stripped:
                break
            if RE_LETTERED.match(lines[i + 1]):
                break
            content += " " + next_stripped
            i += 1

        _add_body(doc, content, indent=True)
        i += 1

    return section_count


def _extract_legal_text(template_content: str) -> str:
    """Extract the legal document text from within ``` code blocks."""
    blocks = re.findall(r"```\n(.*?)```", template_content, re.DOTALL)
    if blocks:
        return blocks[0]
    # If no code blocks, return everything
    return template_content


def _fill_placeholders(text: str, case_data: dict) -> str:
    """Replace {{PLACEHOLDER}} tokens with case data values."""
    parties = case_data.get("parties", {})
    court = case_data.get("court", {})
    plaintiffs = parties.get("plaintiffs", [])
    defendants = parties.get("defendants", [])

    replacements = {
        "PLAINTIFF_NAME": plaintiffs[0]["name"].upper() if plaintiffs else "[PLAINTIFF]",
        "DEFENDANT_NAME": defendants[0]["name"].upper() if defendants else "[DEFENDANT]",
        "CASE_NO": case_data.get("case_number", "________"),
        "ATTORNEY_NAME": case_data.get("attorney", {}).get("name", "[ATTORNEY NAME]"),
        "BAR_NO": case_data.get("attorney", {}).get("bar_number", "[BAR NO.]"),
        "FIRM_NAME": case_data.get("attorney", {}).get("firm", "[FIRM NAME]"),
        "ADDRESS": case_data.get("attorney", {}).get("address", "[ADDRESS]"),
        "CITY": case_data.get("attorney", {}).get("city", "[CITY]"),
        "STATE": case_data.get("attorney", {}).get("state", "[STATE]"),
        "ZIP": case_data.get("attorney", {}).get("zip", "[ZIP]"),
        "PHONE": case_data.get("attorney", {}).get("phone", "[PHONE]"),
        "EMAIL": case_data.get("attorney", {}).get("email", "[EMAIL]"),
        "DATE": case_data.get("filing_date", "____________"),
    }

    for key, value in replacements.items():
        text = text.replace(f"{{{{{key}}}}}", str(value))

    return text


def export_draft(case_data: dict, output_path: str) -> ExportResult:
    """Export a generated complaint draft as .docx."""
    from .drafter import generate_complaint

    complaint_md = generate_complaint(case_data)
    doc = _setup_document()
    sections = _parse_and_render(doc, complaint_md)
    doc.save(output_path)

    return ExportResult(
        output_path=output_path,
        format="docx",
        pages_estimate=max(1, sections),
        sections=sections,
    )


_TEMPLATE_NAME_RE = re.compile(r"^[A-Za-z0-9_\-]+(/[A-Za-z0-9_\-]+)?$")


def export_template(template_name: str, case_data: dict, output_path: str) -> ExportResult:
    """Export a filled template as .docx.

    template_name: e.g. "motions/motion_to_dismiss" or "orders/proposed_order_mtd".
    Security: template_name is restricted to alphanumerics + single subdir, and
    the resolved path is validated to stay under the templates directory.
    """
    if not _TEMPLATE_NAME_RE.match(template_name or ""):
        raise ValueError(
            f"Invalid template name: {template_name!r}. "
            "Allowed: alphanumerics, '_', '-', and a single '/' separator."
        )

    templates_dir = (Path(__file__).parent.parent.parent / "assets" / "templates").resolve()
    candidate = (templates_dir / f"{template_name}.md").resolve()
    if templates_dir not in candidate.parents:
        raise ValueError(f"Template path escapes templates dir: {template_name}")

    template_path = candidate if candidate.exists() else None
    if template_path is None:
        # Fall back: search one level of subfolders for a bare name match
        bare = Path(template_name).name
        for subdir in templates_dir.iterdir():
            if subdir.is_dir():
                c = (subdir / f"{bare}.md").resolve()
                if templates_dir in c.parents and c.exists():
                    template_path = c
                    break
        if template_path is None:
            raise FileNotFoundError(f"Template not found: {template_name}")

    raw = template_path.read_text()
    legal_text = _extract_legal_text(raw)
    filled = _fill_placeholders(legal_text, case_data)

    doc = _setup_document()
    sections = _parse_and_render(doc, filled)
    doc.save(output_path)

    return ExportResult(
        output_path=output_path,
        format="docx",
        pages_estimate=max(1, sections),
        sections=sections,
    )


def export_text(text: str, output_path: str) -> ExportResult:
    """Export raw text/markdown as court-formatted .docx."""
    # Check if it contains code blocks (template format)
    legal_text = _extract_legal_text(text)

    doc = _setup_document()
    sections = _parse_and_render(doc, legal_text)
    doc.save(output_path)

    return ExportResult(
        output_path=output_path,
        format="docx",
        pages_estimate=max(1, sections),
        sections=sections,
    )


def list_templates() -> list[dict]:
    """List all available templates."""
    templates_dir = Path(__file__).parent.parent.parent / "assets" / "templates"
    result = []
    if templates_dir.exists():
        for subdir in sorted(templates_dir.iterdir()):
            if subdir.is_dir():
                for f in sorted(subdir.glob("*.md")):
                    result.append({
                        "category": subdir.name,
                        "name": f.stem,
                        "path": f"{subdir.name}/{f.stem}",
                    })
    return result
