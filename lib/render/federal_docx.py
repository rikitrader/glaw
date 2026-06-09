#!/usr/bin/env python3
"""federal_docx — generate a DOCX that conforms to the GLAW Federal Filing Style Directive
(lib/style/federal-filing-style.md): Times New Roman 12pt body, DOUBLE spacing, JUSTIFIED, 0.5"
first-line indent, margins 1" / 1" / left 1.25" / 1", centered page numbers in the footer, and bold
ALL-CAPS headings/caption. The formatting is APPLIED automatically — callers never hand-set fonts.

Input: a markdown/plain-text filing (file path or stdin). Headings = markdown '#' lines OR ALL-CAPS
lines (caption + I.–VI. section headers); everything else is a justified body paragraph.
Output: a .docx. Verify it with federal_format_check.py.
"""
from __future__ import annotations
import sys, argparse
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Times New Roman"

def _normal_style(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(12)
    rpr = st.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(a), FONT)
    pf = st.paragraph_format
    pf.line_spacing = 2.0            # double
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Inches(0.5)
    pf.space_after = Pt(0)

def _margins(doc):
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1.25)   # binder-friendly
        sec.right_margin = Inches(1)

def _page_numbers(doc):
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for typ, txt in (("begin", None), (None, "PAGE"), ("end", None)):
        if typ:
            f = OxmlElement("w:fldChar"); f.set(qn("w:fldCharType"), typ); run._r.append(f)
        else:
            it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = txt
            run._r.append(it)

def _is_heading(s):
    if s.startswith("#"):
        return True
    letters = [c for c in s if c.isalpha()]
    return bool(letters) and not any(c.islower() for c in s)

def build(text, out):
    doc = Document()
    _normal_style(doc)
    _margins(doc)
    _page_numbers(doc)
    for raw in text.splitlines():
        s = raw.strip()
        if not s:
            continue
        if _is_heading(s):
            head = s.lstrip("#").strip().upper()
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.alignment = (WD_ALIGN_PARAGRAPH.CENTER
                                            if "UNITED STATES DISTRICT COURT" in head
                                            else WD_ALIGN_PARAGRAPH.LEFT)
            r = p.add_run(head); r.bold = True
        else:
            doc.add_paragraph(s)
    doc.save(out)
    return out

def main():
    ap = argparse.ArgumentParser(prog="glaw-federal-format")
    ap.add_argument("path", nargs="?", default="-")
    ap.add_argument("-o", "--out", required=True)
    a = ap.parse_args()
    text = (open(a.path, encoding="utf-8", errors="ignore").read()
            if a.path and a.path != "-" else sys.stdin.read())
    build(text or "FILING", a.out)
    print(f"federal DOCX written: {a.out}  (TNR 12pt · double · justified · 1.25\" left margin · page #s)")
    return 0

if __name__ == "__main__":
    raise SystemExit(main())
