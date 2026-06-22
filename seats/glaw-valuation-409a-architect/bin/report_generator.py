#!/usr/bin/env python3
"""
409A Valuation Report Generator.

Turns the engine's `results.json` + `audit_log.json` into the 12-section
Markdown report (and a DOCX when python-docx is installed). Reproducible — the
report is rendered from data, never hand-written.

Usage:
    python3 report_generator.py --results out/results.json --audit out/audit_log.json \
        [--intake intake.json] [--out report.md] [--docx report.docx]

If --docx is given and python-docx is unavailable, it degrades to Markdown only
and says so. Mirrors references/report-template.md.
"""
from __future__ import annotations

import argparse
import json


def m(x):
    return f"${x:,.2f}" if isinstance(x, (int, float)) else str(x)


def pct(x):
    return f"{x*100:.1f}%" if isinstance(x, (int, float)) else str(x)


DISCLAIMER = (
    "> **DRAFT — not a safe-harbor appraisal.** This report was produced by the "
    "glaw-valuation-409a-architect engine. It is **not** an independent appraisal "
    "and does not by itself establish the IRC section 409A safe-harbor presumption, "
    "which requires qualified independent appraiser review and sign-off. "
    "Attorney/CPA work-product for licensed professional review; not legal, tax, "
    "accounting, or appraisal advice."
)


def build_markdown(results, audit, intake=None):
    r = results
    head = r.get("headline", {})
    ent = (intake or {}).get("entity", {}) if intake else {}
    L = []
    A = L.append

    A(f"# 409A Valuation Report — {ent.get('company_name', '(company)')}")
    A("")
    A(DISCLAIMER)
    A("")
    A(f"*Valuation date: {audit.get('valuation_date','?')} · Engine v{audit.get('engine_version','?')} "
      f"· Generated {audit.get('generated_at','?')}*")
    A("")

    # 1. Executive Summary
    A("## 1. Executive Summary")
    A("")
    A("| Metric | Value |")
    A("|---|---|")
    A(f"| Enterprise Value | {m(head.get('enterprise_value'))} |")
    A(f"| Equity Value | {m(head.get('equity_value'))} |")
    A(f"| Common FMV (marketable, waterfall) | {m(head.get('common_fmv_marketable'))} |")
    if head.get("opm_fmv_marketable") is not None:
        A(f"| Common FMV (marketable, OPM) | {m(head.get('opm_fmv_marketable'))} |")
    A(f"| DLOM | {pct(head.get('dlom'))} |")
    A(f"| **Recommended 409A Strike** | **{m(head.get('recommended_409a_strike'))}** |")
    comp = r.get("compliance", {})
    if comp:
        A(f"| Compliance | {comp.get('status')} |")
    legal = r.get("legal_audit", {})
    if legal:
        A(f"| Legal/Appraiser Audit Gate | {legal.get('status')} |")
    A("")

    # 2. Company Overview
    A("## 2. Company Overview")
    A("")
    if ent:
        A(f"- **Entity:** {ent.get('company_name','?')} ({ent.get('state_of_incorporation','?')})")
        A(f"- **Industry:** {ent.get('industry','?')}")
    for rnd in (intake or {}).get("financing_rounds", []) if intake else []:
        A(f"- **Round:** {rnd.get('round')} — {m(rnd.get('amount'))} @ "
          f"{m(rnd.get('price_per_share'))}/sh ({rnd.get('date')})")
    A("")

    # 3. Financial Analysis
    A("## 3. Financial Analysis")
    A("")
    fin = (intake or {}).get("financials", {}) if intake else {}
    if fin:
        A(f"- Cash {m(fin.get('cash'))} · Debt {m(fin.get('debt'))} · "
          f"Net debt {m((fin.get('debt',0) or 0)-(fin.get('cash',0) or 0))}")
    rec = r.get("reconciliation", {})
    if rec:
        A(f"- Approach EVs: {rec.get('approach_evs')}")
        A(f"- Weights used: {rec.get('weights_used')}")
    round_basis = r.get("priced_round")
    if round_basis:
        A(f"- Latest priced round anchor: {round_basis.get('round')} "
          f"({round_basis.get('basis')}) -> equity value {m(round_basis.get('equity_value'))}")
    support = r.get("valuation_support") or {}
    if support.get("approach_dispersion") is not None:
        A(f"- Approach dispersion: {pct(support.get('approach_dispersion'))}")
    A("")

    # 4. Market Approach
    A("## 4. Market Approach (Comparables)")
    A("")
    comps = r.get("comps")
    if comps:
        for rm in comps.get("recommended_multiples", []):
            A(f"- {rm['metric']}: {rm['multiple']}× → implied EV {m(rm['implied_ev'])}")
        A(f"- **Comps enterprise value:** {m(comps.get('enterprise_value'))}")
    else:
        A("- No comparable set provided; market approach omitted.")
    A("")

    # 5 + 6. Income / DCF
    A("## 5–6. Income Approach (DCF)")
    A("")
    dcf = r.get("dcf")
    if dcf:
        A("| Year | Revenue | FCF | PV |")
        A("|---|---|---|---|")
        for y in dcf.get("years", []):
            A(f"| {y['year']} | {m(y['revenue'])} | {m(y['fcf'])} | {m(y['pv'])} |")
        A(f"- Terminal value {m(dcf.get('terminal_value'))} → PV {m(dcf.get('pv_terminal_value'))}")
        ev = dcf.get("enterprise_value") or 0
        pvtv = dcf.get("pv_terminal_value") or 0
        share = (pvtv / ev) if ev else 0
        A(f"- **DCF enterprise value:** {m(ev)} (PV of terminal value = {pct(share)} of EV"
          + (" — review" if share > 0.75 else "") + ")")
    else:
        A("- No forecast provided; income approach omitted.")
    A("")

    # 7. PWERM
    A("## 7. PWERM Analysis")
    A("")
    pw = r.get("pwerm")
    if pw:
        A("| Scenario | Prob | Exit | t | PV | Weighted |")
        A("|---|---|---|---|---|---|")
        for s in pw.get("scenarios", []):
            A(f"| {s['scenario']} | {pct(s['probability'])} | {m(s['exit_value'])} | "
              f"{s['time_to_exit']} | {m(s['present_value'])} | {m(s['weighted_contribution'])} |")
        A(f"- **Probability-weighted equity value:** {m(pw.get('probability_weighted_equity_value'))}")
        if rec and rec.get("pwerm_divergence") is not None:
            A(f"- Divergence vs reconciled equity value: {pct(rec['pwerm_divergence'])}")
    else:
        A("- No scenarios provided; PWERM omitted.")
    A("")

    # 8. Discount (DLOM)
    A("## 8. Discount Analysis (DLOM)")
    A("")
    dl = r.get("dlom")
    if dl:
        A(f"- Score {dl.get('score')} → raw {pct(dl.get('raw_dlom'))} → "
          f"**recommended DLOM {pct(dl.get('recommended_dlom'))}**")
    sens = r.get("sensitivity")
    if sens:
        A("")
        A("### Sensitivity")
        A("")
        A("| Equity case | Equity factor | DLOM case | DLOM | Strike |")
        A("|---|---:|---|---:|---:|")
        for row in sens.get("rows", []):
            A(f"| {row['equity_case']} | {row['equity_factor']} | {row['dlom_case']} | "
              f"{pct(row['dlom'])} | {m(row['strike'])} |")
    A("")

    # 9. FMV Conclusion
    A("## 9. FMV Conclusion")
    A("")
    wf = r.get("waterfall")
    if wf:
        if wf.get("converted_to_common"):
            A(f"- Preferred converting to common: {', '.join(wf['converted_to_common'])}")
        A(f"- Preferences paid {m(wf.get('preferred_preferences_paid'))}; "
          f"participation {m(wf.get('participation_paid'))}; "
          f"**value to common {m(wf.get('value_to_common'))}**")
    fmv = r.get("common_fmv")
    sp = r.get("strike")
    if fmv and sp:
        A(f"- Common FMV (marketable) {m(fmv.get('fmv_marketable'))} × (1 − "
          f"{pct(sp.get('dlom'))}) = **strike {m(sp.get('strike_price'))}**")
    opm = r.get("opm")
    if opm:
        A(f"- OPM cross-check FMV {m(opm.get('fmv_marketable'))} "
          f"(σ {opm.get('sigma')}, T {opm.get('years')}y); "
          f"divergence vs waterfall {pct(opm.get('divergence_vs_waterfall', 0))}")
    bs = r.get("opm_backsolve")
    if bs:
        A(f"- Backsolve from {bs.get('round')}: equity {m(bs.get('backsolved_equity_value'))}")
    A("")

    # 9b. Reviewer support
    if support:
        A("### Reviewer Support Workpaper")
        A("")
        bt = support.get("backsolve_tieout") or {}
        if bt:
            A(f"- Backsolve tie-out: {bt.get('round')} solved equity {m(bt.get('backsolved_equity_value'))}; "
              f"round post-money {m(bt.get('implied_post_money'))}; divergence "
              f"{pct(bt.get('divergence_vs_round_post_money'))}.")
        vol = support.get("volatility_benchmark") or {}
        if vol:
            A(f"- Volatility benchmark: sigma {vol.get('sigma')} -> {vol.get('review_band')}. {vol.get('note')}")
        dlom_support = support.get("dlom_support") or {}
        if dlom_support:
            A(f"- DLOM support: {pct(dlom_support.get('recommended_dlom'))} -> "
              f"{dlom_support.get('review_band')}. {dlom_support.get('note')}")
        comps_support = support.get("comps_scoring") or {}
        if comps_support:
            A(f"- Comps support: {comps_support.get('peer_count')} peers; EV/Revenue range/median "
              f"{comps_support.get('ev_revenue_range_over_median')}. {comps_support.get('note')}")
        if support.get("pwerm_sensitivity"):
            A("")
            A("| PWERM case | Factor | Equity value |")
            A("|---|---:|---:|")
            for row in support.get("pwerm_sensitivity", []):
                A(f"| {row['case']} | {row['factor']} | {m(row['equity_value'])} |")
        A("")

    # 10. Assumptions
    A("## 10. Assumptions")
    A("")
    A("All parameters, including engine-applied defaults, are recorded step-by-step "
      "in the audit trail (Appendix B). Key drivers:")
    if dcf:
        A(f"- Discount rate {pct(dcf.get('discount_rate'))}, terminal growth {pct(dcf.get('terminal_growth'))}")
    if opm:
        A(f"- OPM volatility {opm.get('sigma')}, time-to-liquidity {opm.get('years')}y, risk-free {pct(opm.get('rate'))}")
    A("")

    # 11. Limitations
    A("## 11. Limitations")
    A("")
    A(DISCLAIMER)
    A("")
    for w in audit.get("warnings", []):
        A(f"- WARNING: {w}")
    A("")

    # 12. Compliance
    A("## 12. Compliance Statement")
    A("")
    if comp:
        A(f"- **Status: {comp.get('status')}**")
        for t in comp.get("triggers", []):
            A(f"  - Trigger: {t}")
    A("- Re-review within 12 months or on any material event.")
    if legal:
        A("")
        A("### Legal / Appraiser Audit Gate")
        A("")
        A("- Checklist basis: `references/skadden-409a-equity-pitfalls.md` "
          "(Skadden-hosted Practical Law 409A equity-award pitfalls checklist).")
        A(f"- **Status: {legal.get('status')}**")
        for item in legal.get("missing_controls", []):
            A(f"- Open control: {item}")
        if not legal.get("missing_controls"):
            A("- No open review controls in the intake checklist.")
    A("")

    # Appendix B — Audit trail
    A("## Appendix B — Audit Trail")
    A("")
    A("| # | Module | Formula | Result |")
    A("|---|---|---|---|")
    for i, step in enumerate(audit.get("steps", []), 1):
        res = json.dumps(step.get("result", {}))
        A(f"| {i} | {step['module']} | {step['formula'][:80]} | {res[:80]} |")
    A("")
    A("## Appendix C — Adversarial RED/BLUE Review")
    A("")
    A("Fill from `references/seats-and-adversary.md` against the numbers above "
      "(IRS section 409A examiner, safe-harbor reliance, cheap-stock, ASC 718/820 "
      "auditor, valuation analyst, board, CFO, tax counsel). Each lens -> HOLDS / "
      "HOLDS-CONDITIONED / MATERIAL. The safe-harbor reliance lens remains MATERIAL "
      "until a qualified independent appraiser signs.")
    A("")
    A("## Appendix D — Required Workpaper Templates")
    A("")
    A("- `templates/board-approval-checklist.md`")
    A("- `templates/appraiser-signoff-checklist.md`")
    A("- `templates/asc718-820-audit-workpaper.md`")
    A("- `templates/irs-409a-legal-risk-matrix.md`")
    A("")
    return "\n".join(L)


def to_docx(md_text, path):
    try:
        import docx  # noqa: F401
        from docx import Document
    except ImportError:
        return False
    doc = Document()
    for line in md_text.split("\n"):
        if line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("|"):
            doc.add_paragraph(line, style="Intense Quote")
        elif line.startswith("- ") or line.startswith("  - "):
            doc.add_paragraph(line.lstrip("- "), style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line)
    doc.save(path)
    return True


def main():
    ap = argparse.ArgumentParser(description="409A report generator")
    ap.add_argument("--results", required=True)
    ap.add_argument("--audit", required=True)
    ap.add_argument("--intake", default=None)
    ap.add_argument("--out", default="report.md")
    ap.add_argument("--docx", default=None)
    args = ap.parse_args()

    results = json.load(open(args.results))
    audit = json.load(open(args.audit))
    intake = json.load(open(args.intake)) if args.intake else None

    md = build_markdown(results, audit, intake)
    with open(args.out, "w") as f:
        f.write(md)
    print(f"Wrote {args.out}")

    if args.docx:
        if to_docx(md, args.docx):
            print(f"Wrote {args.docx}")
        else:
            print("python-docx not installed — DOCX skipped, Markdown delivered.")


if __name__ == "__main__":
    main()
